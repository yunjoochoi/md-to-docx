from docx import Document
from markdown_it import MarkdownIt

def md_to_word_with_style(md_text, template_path, output_path):
    # 1. 템플릿 로드 (Pandoc의 reference-doc 역할)
    doc = Document(template_path)
    
    # 본문 초기화 (스타일만 남기고 내용 삭제)
    for element in doc.element.body:
        doc.element.body.remove(element)

    # 2. 마크다운 파서 준비
    md = MarkdownIt()
    tokens = md.parse(md_text) # 마크다운을 토큰 리스트로 변환

    # 3. 토큰을 순회하며 워드 스타일 매핑
    for i, token in enumerate(tokens):
        
        # 제목 (Heading) 처리
        if token.type == 'heading_open':
            # h1 -> Heading 1, h2 -> Heading 2 ...
            level = token.tag  # 'h1', 'h2'
            style_name = f"Heading {level[1]}" 
            
            # 다음 토큰(실제 텍스트)을 가져와서 워드에 삽입
            content = tokens[i+1].content 
            doc.add_paragraph(content, style=style_name)

        # 본문 (Paragraph) 처리
        elif token.type == 'paragraph_open':
            # 리스트 내부인지 아닌지 등을 판단하여 스타일 적용
            # 단순 본문이라면:
            if tokens[i+1].type == 'inline':
                content = tokens[i+1].content
                doc.add_paragraph(content, style='Normal')

        # 리스트 아이템 처리
        elif token.type == 'list_item_open':
             # 리스트 내부 텍스트 처리는 로직이 조금 더 필요함
             pass

    # 4. 저장
    doc.save(output_path)