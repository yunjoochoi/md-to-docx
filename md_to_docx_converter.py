"""
마크다운 → DOCX 변환기

- markdown-it 파서를 사용한 토큰 기반 변환
- 템플릿 docx에서 추출한 스타일 적용
- 배경 이미지 유지
- 하드코딩된 텍스트 지표 없이 마크다운 문법만으로 매핑
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from markdown_it import MarkdownIt
from markdown_it.token import Token
from pathlib import Path
from typing import List, Dict, Optional
import re
import copy


class MarkdownToDocxConverter:
    """마크다운을 DOCX로 변환하는 클래스"""

    # 마크다운 토큰 → DOCX 스타일 매핑
    DEFAULT_STYLE_MAP = {
        'heading_1': 'Heading 1',
        'heading_2': 'Heading 2',
        'heading_3': 'Heading 3',
        'heading_4': 'Heading 4',
        'heading_5': 'Heading 5',
        'heading_6': 'Heading 6',
        'paragraph': 'Normal',
        'bullet_list': 'List Bullet',
        'ordered_list': 'List Number',
        'blockquote': 'Quote',
        'code_block': 'Normal',
        'table': 'Table Grid',
    }

    def __init__(self, template_path: Optional[str] = None):
        """
        Args:
            template_path: DOCX 템플릿 파일 경로 (None이면 빈 문서 생성)
        """
        self.template_path = template_path
        self.md_parser = MarkdownIt('commonmark', {'breaks': True, 'html': True})
        self.md_parser.enable('table')  # 테이블 지원

        # 템플릿 로드 또는 새 문서 생성
        if template_path and Path(template_path).exists():
            self.doc = Document(template_path)
        else:
            self.doc = Document()

        self.style_map = self.DEFAULT_STYLE_MAP.copy()

    def set_style_map(self, custom_map: Dict[str, str]):
        """커스텀 스타일 매핑 설정"""
        self.style_map.update(custom_map)

    def convert(self, md_content: str, output_path: str, clear_body: bool = True):
        """
        마크다운을 DOCX로 변환

        Args:
            md_content: 마크다운 텍스트
            output_path: 출력 DOCX 파일 경로
            clear_body: True면 템플릿의 본문 내용을 지우고 시작
        """
        if clear_body:
            self._clear_body()

        # 마크다운 파싱
        tokens = self.md_parser.parse(md_content)

        # 토큰 처리
        self._process_tokens(tokens)

        # 저장
        self.doc.save(output_path)
        return output_path

    def convert_file(self, md_file_path: str, output_path: str, clear_body: bool = True):
        """마크다운 파일을 DOCX로 변환"""
        with open(md_file_path, 'r', encoding='utf-8') as f:
            md_content = f.read()
        return self.convert(md_content, output_path, clear_body)

    def _clear_body(self):
        """문서 본문 내용 삭제 (헤더/푸터/스타일은 유지)"""
        # 본문 요소만 삭제 (섹션 속성 유지)
        body = self.doc.element.body
        for child in list(body):
            # sectPr (섹션 속성)은 유지
            if child.tag.endswith('sectPr'):
                continue
            body.remove(child)

    def _process_tokens(self, tokens: List[Token]):
        """토큰 리스트 처리"""
        i = 0
        while i < len(tokens):
            token = tokens[i]
            consumed = self._handle_token(token, tokens, i)
            i += consumed

    def _handle_token(self, token: Token, tokens: List[Token], index: int) -> int:
        """
        단일 토큰 처리

        Returns:
            처리한 토큰 수 (다음 인덱스로 건너뛸 양)
        """
        token_type = token.type

        # 헤딩 처리 (h1~h6)
        if token_type == 'heading_open':
            return self._handle_heading(tokens, index)

        # 문단 처리
        elif token_type == 'paragraph_open':
            return self._handle_paragraph(tokens, index)

        # 불릿 리스트
        elif token_type == 'bullet_list_open':
            return self._handle_bullet_list(tokens, index)

        # 숫자 리스트
        elif token_type == 'ordered_list_open':
            return self._handle_ordered_list(tokens, index)

        # 인용문
        elif token_type == 'blockquote_open':
            return self._handle_blockquote(tokens, index)

        # 코드 블록
        elif token_type == 'fence' or token_type == 'code_block':
            return self._handle_code_block(token)

        # 테이블
        elif token_type == 'table_open':
            return self._handle_table(tokens, index)

        # 수평선
        elif token_type == 'hr':
            self._add_horizontal_rule()
            return 1

        # 이미지 (paragraph 안에 있음)
        elif token_type == 'image':
            self._handle_image(token)
            return 1

        # 기타 (무시)
        return 1

    def _handle_heading(self, tokens: List[Token], start: int) -> int:
        """헤딩 처리 (heading_open, inline, heading_close)"""
        open_token = tokens[start]
        level = int(open_token.tag[1])  # h1 -> 1

        # inline 토큰에서 텍스트 추출
        inline_token = tokens[start + 1]
        text = self._extract_text_from_inline(inline_token)

        # 스타일 적용
        style_key = f'heading_{level}'
        style_name = self.style_map.get(style_key, f'Heading {level}')

        try:
            para = self.doc.add_paragraph(style=style_name)
        except KeyError:
            # 스타일이 없으면 Normal로 대체하고 수동 서식
            para = self.doc.add_paragraph(style='Normal')

        self._add_formatted_text(para, inline_token)

        return 3  # open, inline, close

    def _handle_paragraph(self, tokens: List[Token], start: int) -> int:
        """문단 처리"""
        if start + 1 >= len(tokens):
            return 1

        inline_token = tokens[start + 1]
        if inline_token.type != 'inline':
            return 1

        # 이미지만 있는 경우 특별 처리
        if inline_token.children:
            is_only_image = all(
                child.type in ('image', 'softbreak', 'hardbreak')
                for child in inline_token.children
            )
            if is_only_image:
                for child in inline_token.children:
                    if child.type == 'image':
                        self._handle_image(child)
                return 3

        # 일반 문단
        style_name = self.style_map.get('paragraph', 'Normal')
        try:
            para = self.doc.add_paragraph(style=style_name)
        except KeyError:
            para = self.doc.add_paragraph()

        self._add_formatted_text(para, inline_token)

        return 3  # open, inline, close

    def _handle_bullet_list(self, tokens: List[Token], start: int) -> int:
        """불릿 리스트 처리"""
        consumed = 1
        depth = 1

        i = start + 1
        while i < len(tokens):
            token = tokens[i]

            if token.type == 'bullet_list_close':
                depth -= 1
                if depth == 0:
                    consumed = i - start + 1
                    break

            elif token.type == 'bullet_list_open':
                depth += 1

            elif token.type == 'list_item_open':
                pass

            elif token.type == 'paragraph_open':
                if i + 1 < len(tokens) and tokens[i + 1].type == 'inline':
                    inline_token = tokens[i + 1]
                    style_name = self.style_map.get('bullet_list', 'List Bullet')
                    try:
                        para = self.doc.add_paragraph(style=style_name)
                    except KeyError:
                        para = self.doc.add_paragraph()
                        para.text = '• '  # 수동 불릿
                    self._add_formatted_text(para, inline_token)
                    i += 2  # paragraph_open, inline 건너뛰기
                    continue

            i += 1

        return consumed

    def _handle_ordered_list(self, tokens: List[Token], start: int) -> int:
        """숫자 리스트 처리"""
        consumed = 1
        depth = 1
        item_num = 1

        i = start + 1
        while i < len(tokens):
            token = tokens[i]

            if token.type == 'ordered_list_close':
                depth -= 1
                if depth == 0:
                    consumed = i - start + 1
                    break

            elif token.type == 'ordered_list_open':
                depth += 1

            elif token.type == 'list_item_open':
                pass

            elif token.type == 'paragraph_open':
                if i + 1 < len(tokens) and tokens[i + 1].type == 'inline':
                    inline_token = tokens[i + 1]
                    style_name = self.style_map.get('ordered_list', 'List Number')
                    try:
                        para = self.doc.add_paragraph(style=style_name)
                    except KeyError:
                        para = self.doc.add_paragraph()
                        # 수동 번호 추가는 하지 않음 (스타일에 의존)
                    self._add_formatted_text(para, inline_token)
                    item_num += 1
                    i += 2
                    continue

            i += 1

        return consumed

    def _handle_blockquote(self, tokens: List[Token], start: int) -> int:
        """인용문 처리"""
        consumed = 1
        depth = 1

        i = start + 1
        while i < len(tokens):
            token = tokens[i]

            if token.type == 'blockquote_close':
                depth -= 1
                if depth == 0:
                    consumed = i - start + 1
                    break

            elif token.type == 'blockquote_open':
                depth += 1

            elif token.type == 'paragraph_open':
                if i + 1 < len(tokens) and tokens[i + 1].type == 'inline':
                    inline_token = tokens[i + 1]
                    style_name = self.style_map.get('blockquote', 'Quote')
                    try:
                        para = self.doc.add_paragraph(style=style_name)
                    except KeyError:
                        para = self.doc.add_paragraph()
                        # 인용문 들여쓰기
                        para.paragraph_format.left_indent = Inches(0.5)
                    self._add_formatted_text(para, inline_token)
                    i += 2
                    continue

            i += 1

        return consumed

    def _handle_code_block(self, token: Token) -> int:
        """코드 블록 처리"""
        code_text = token.content.rstrip('\n')
        try:
            para = self.doc.add_paragraph(style='Normal')
        except:
            para = self.doc.add_paragraph()

        # 코드 스타일 적용
        run = para.add_run(code_text)
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        # 배경색은 python-docx로 직접 설정하기 어려움

        return 1

    def _handle_table(self, tokens: List[Token], start: int) -> int:
        """테이블 처리"""
        # 테이블 구조 파싱
        rows = []
        current_row = []
        in_header = False
        consumed = 1

        i = start + 1
        while i < len(tokens):
            token = tokens[i]

            if token.type == 'table_close':
                consumed = i - start + 1
                break

            elif token.type == 'thead_open':
                in_header = True

            elif token.type == 'tr_open':
                current_row = []

            elif token.type == 'tr_close':
                rows.append((current_row, in_header))

            elif token.type == 'thead_close':
                in_header = False

            elif token.type in ('th_open', 'td_open'):
                pass

            elif token.type == 'inline':
                cell_text = self._extract_text_from_inline(token)
                current_row.append(cell_text)

            i += 1

        # DOCX 테이블 생성
        if rows:
            num_cols = max(len(row[0]) for row in rows) if rows else 0
            table = self.doc.add_table(rows=len(rows), cols=num_cols)
            table.style = 'Table Grid'

            for row_idx, (cells, is_header) in enumerate(rows):
                for col_idx, cell_text in enumerate(cells):
                    if col_idx < num_cols:
                        cell = table.rows[row_idx].cells[col_idx]
                        cell.text = cell_text
                        if is_header:
                            for para in cell.paragraphs:
                                for run in para.runs:
                                    run.bold = True

        return consumed

    def _handle_image(self, token: Token):
        """이미지 처리"""
        src = token.attrGet('src') or ''
        alt = token.attrGet('alt') or token.content or ''

        # 이미지 경로가 상대 경로인 경우 처리
        # 실제 이미지 삽입은 이미지 파일이 존재해야 함
        para = self.doc.add_paragraph()
        if Path(src).exists():
            try:
                para.add_run().add_picture(src, width=Inches(5))
            except Exception as e:
                para.add_run(f"[이미지: {alt or src}]")
        else:
            para.add_run(f"[이미지: {alt or src}]")

    def _add_horizontal_rule(self):
        """수평선 추가"""
        para = self.doc.add_paragraph()
        para.add_run('─' * 50)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _extract_text_from_inline(self, inline_token: Token) -> str:
        """inline 토큰에서 순수 텍스트 추출"""
        if not inline_token.children:
            return inline_token.content or ''

        text_parts = []
        for child in inline_token.children:
            if child.type == 'text':
                text_parts.append(child.content)
            elif child.type == 'code_inline':
                text_parts.append(child.content)
            elif child.type == 'softbreak':
                text_parts.append(' ')
            elif child.type == 'hardbreak':
                text_parts.append('\n')
            elif child.children:
                text_parts.append(self._extract_text_from_inline(child))

        return ''.join(text_parts)

    def _add_formatted_text(self, para, inline_token: Token):
        """인라인 토큰의 서식을 적용하며 텍스트 추가"""
        if not inline_token.children:
            if inline_token.content:
                para.add_run(inline_token.content)
            return

        self._process_inline_children(para, inline_token.children, {})

    def _process_inline_children(self, para, children: List[Token], format_state: dict):
        """인라인 자식 토큰 처리 (재귀)"""
        for child in children:
            if child.type == 'text':
                run = para.add_run(child.content)
                self._apply_format(run, format_state)

            elif child.type == 'code_inline':
                run = para.add_run(child.content)
                run.font.name = 'Consolas'
                run.font.size = Pt(10)

            elif child.type == 'softbreak':
                para.add_run(' ')

            elif child.type == 'hardbreak':
                para.add_run('\n')

            elif child.type == 'strong_open':
                new_state = format_state.copy()
                new_state['bold'] = True
                # 다음 토큰들 처리
                idx = children.index(child)
                inner_children = []
                for j in range(idx + 1, len(children)):
                    if children[j].type == 'strong_close':
                        break
                    inner_children.append(children[j])
                self._process_inline_children(para, inner_children, new_state)

            elif child.type == 'em_open':
                new_state = format_state.copy()
                new_state['italic'] = True
                idx = children.index(child)
                inner_children = []
                for j in range(idx + 1, len(children)):
                    if children[j].type == 'em_close':
                        break
                    inner_children.append(children[j])
                self._process_inline_children(para, inner_children, new_state)

            elif child.type == 'link_open':
                # 링크는 텍스트만 추출
                pass

            elif child.type in ('strong_close', 'em_close', 'link_close'):
                pass  # 이미 처리됨

            elif child.type == 's_open':  # 취소선
                new_state = format_state.copy()
                new_state['strike'] = True
                idx = children.index(child)
                inner_children = []
                for j in range(idx + 1, len(children)):
                    if children[j].type == 's_close':
                        break
                    inner_children.append(children[j])
                self._process_inline_children(para, inner_children, new_state)

    def _apply_format(self, run, format_state: dict):
        """run에 서식 적용"""
        if format_state.get('bold'):
            run.bold = True
        if format_state.get('italic'):
            run.italic = True
        if format_state.get('strike'):
            run.font.strike = True


def main():
    """테스트 실행"""
    import sys

    # 테스트 마크다운
    test_md = """
# 제목 1

이것은 **굵은 텍스트**와 *기울임 텍스트*가 포함된 문단입니다.

## 제목 2

- 불릿 항목 1
- 불릿 항목 2
- 불릿 항목 3

### 제목 3

1. 숫자 항목 1
2. 숫자 항목 2
3. 숫자 항목 3

> 이것은 인용문입니다.
> 여러 줄로 작성할 수 있습니다.

#### 테이블 예시

| 항목 | 설명 | 비고 |
|------|------|------|
| A | 첫번째 | 테스트 |
| B | 두번째 | 데이터 |

```python
def hello():
    print("Hello, World!")
```

---

일반 문단으로 마무리합니다.
"""

    template_path = '/home/shaush/md-to-docx/docx_only/[Word 템플릿] A4.docx'
    output_path = '/home/shaush/md-to-docx/test_output.docx'

    if len(sys.argv) > 1:
        # 실제 마크다운 파일 변환
        md_file = sys.argv[1]
        with open(md_file, 'r', encoding='utf-8') as f:
            test_md = f.read()
        output_path = Path(md_file).stem + '_converted.docx'

    converter = MarkdownToDocxConverter(template_path)
    result = converter.convert(test_md, output_path)
    print(f"✅ 변환 완료: {result}")


if __name__ == '__main__':
    main()
