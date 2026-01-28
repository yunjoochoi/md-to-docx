"""
DOCX 템플릿 플레이스홀더 파서

- {{TITLE}}, {{BODY}} 등 플레이스홀더 패턴 추출
- 각 플레이스홀더의 스타일 정보 저장
- 섹션 (cover, toc, body) 분류
"""

import re
from pathlib import Path
from typing import List, Optional, Dict, Any
from docx import Document
from docx.text.paragraph import Paragraph

from .models import (
    Placeholder, ParsedTemplate, PlaceholderType, SectionType,
    PLACEHOLDER_PATTERNS, parse_placeholder_id
)
from .template_analyzer import DocxTemplateAnalyzer


class TemplateParser:
    """
    DOCX 템플릿에서 플레이스홀더를 추출

    지원 패턴:
    - {{TITLE}}, {{BODY}}, {{SECTION_1}} 등
    - [[TITLE]] (대안 패턴)
    - <<TITLE>> (대안 패턴)
    """

    def __init__(
        self,
        docx_path: str,
        placeholder_pattern: str = "default"
    ):
        """
        Args:
            docx_path: DOCX 템플릿 파일 경로
            placeholder_pattern: 사용할 플레이스홀더 패턴
                - "default": {{TITLE}}
                - "bracket": [[TITLE]]
                - "angle": <<TITLE>>
                - "underscore": ___TITLE___
        """
        self.docx_path = Path(docx_path)
        self.doc = Document(docx_path)
        self.pattern = PLACEHOLDER_PATTERNS.get(placeholder_pattern, PLACEHOLDER_PATTERNS["default"])
        self.regex = re.compile(self.pattern)

        # 템플릿 분석기 (스타일 정보 획득용)
        self._analyzer: Optional[DocxTemplateAnalyzer] = None

    def parse(self) -> ParsedTemplate:
        """
        템플릿 파싱하여 플레이스홀더 추출

        Returns:
            ParsedTemplate 객체
        """
        result = ParsedTemplate(
            file_path=str(self.docx_path),
            total_paragraphs=len(self.doc.paragraphs)
        )

        # 문단별로 플레이스홀더 추출
        for para_idx, paragraph in enumerate(self.doc.paragraphs):
            placeholders = self._extract_from_paragraph(paragraph, para_idx)
            result.placeholders.extend(placeholders)

        # 헤더/푸터에서도 추출
        for section in self.doc.sections:
            # 헤더
            if section.header:
                for para_idx, paragraph in enumerate(section.header.paragraphs):
                    placeholders = self._extract_from_paragraph(
                        paragraph, para_idx, prefix="header_"
                    )
                    result.placeholders.extend(placeholders)

            # 푸터
            if section.footer:
                for para_idx, paragraph in enumerate(section.footer.paragraphs):
                    placeholders = self._extract_from_paragraph(
                        paragraph, para_idx, prefix="footer_"
                    )
                    result.placeholders.extend(placeholders)

        # 섹션별 분류
        self._classify_placeholders(result)

        return result

    def _extract_from_paragraph(
        self,
        paragraph: Paragraph,
        para_idx: int,
        prefix: str = ""
    ) -> List[Placeholder]:
        """문단에서 플레이스홀더 추출"""
        placeholders = []
        text = paragraph.text

        # 정규식으로 모든 플레이스홀더 찾기
        for match in self.regex.finditer(text):
            raw_id = match.group(0)  # "{{TITLE}}"
            clean_id = match.group(1)  # "TITLE"

            # 타입과 섹션 번호 파싱
            ptype, section_num = parse_placeholder_id(raw_id)

            # 스타일 정보 가져오기
            style_id = None
            style_name = None
            if paragraph.style:
                style_id = paragraph.style.style_id
                style_name = paragraph.style.name

            # run 위치 찾기
            run_index = self._find_run_index(paragraph, match.start())

            placeholder = Placeholder(
                id=raw_id,
                placeholder_type=ptype,
                paragraph_index=para_idx,
                run_index=run_index,
                style_id=style_id,
                style_name=style_name,
                section_number=section_num,
                original_text=text
            )

            placeholders.append(placeholder)

        return placeholders

    def _find_run_index(self, paragraph: Paragraph, char_pos: int) -> int:
        """문자 위치로 run 인덱스 찾기"""
        current_pos = 0
        for run_idx, run in enumerate(paragraph.runs):
            run_len = len(run.text)
            if current_pos + run_len > char_pos:
                return run_idx
            current_pos += run_len
        return 0

    def _classify_placeholders(self, result: ParsedTemplate):
        """플레이스홀더를 섹션별로 분류"""
        # 간단한 휴리스틱:
        # - TITLE, SUBTITLE은 보통 cover
        # - TOC는 toc
        # - 나머지는 body

        for placeholder in result.placeholders:
            if placeholder.placeholder_type in [PlaceholderType.TITLE, PlaceholderType.SUBTITLE]:
                placeholder.section_type = SectionType.COVER
                result.cover_placeholders.append(placeholder)
            elif placeholder.placeholder_type == PlaceholderType.TOC:
                placeholder.section_type = SectionType.TOC
                result.toc_placeholders.append(placeholder)
            else:
                placeholder.section_type = SectionType.BODY
                result.body_placeholders.append(placeholder)

    def get_analyzer(self) -> DocxTemplateAnalyzer:
        """템플릿 분석기 인스턴스 반환 (lazy loading)"""
        if self._analyzer is None:
            self._analyzer = DocxTemplateAnalyzer(str(self.docx_path))
            self._analyzer.analyze()
        return self._analyzer

    def get_style_info(self, style_id: str) -> Optional[Dict[str, Any]]:
        """스타일 ID로 스타일 정보 가져오기"""
        analyzer = self.get_analyzer()
        style = analyzer.structure.styles.get(style_id)
        if style:
            return style.to_dict()
        return None

    def to_dict(self) -> Dict[str, Any]:
        """파싱 결과를 딕셔너리로 변환 (LLM 입력용)"""
        result = self.parse()
        return {
            "file_path": result.file_path,
            "total_paragraphs": result.total_paragraphs,
            "placeholders": [
                {
                    "id": p.id,
                    "type": p.placeholder_type,
                    "section": p.section_type,
                    "style": p.style_name,
                    "paragraph_index": p.paragraph_index,
                }
                for p in result.placeholders
            ]
        }


def find_placeholders_in_text(text: str, pattern: str = "default") -> List[str]:
    """
    텍스트에서 플레이스홀더 ID 목록 추출

    Args:
        text: 검색할 텍스트
        pattern: 플레이스홀더 패턴 ("default", "bracket", "angle", "underscore")

    Returns:
        플레이스홀더 ID 목록 (예: ["{{TITLE}}", "{{BODY}}"])
    """
    regex = re.compile(PLACEHOLDER_PATTERNS.get(pattern, PLACEHOLDER_PATTERNS["default"]))
    return [match.group(0) for match in regex.finditer(text)]


if __name__ == "__main__":
    import sys
    import json

    # 테스트 실행
    template_path = "/home/shaush/md-to-docx/[Word템플릿]A4.docx"
    if len(sys.argv) > 1:
        template_path = sys.argv[1]

    parser = TemplateParser(template_path)
    result = parser.parse()

    print(f"\n{'='*60}")
    print(f"템플릿 플레이스홀더 분석: {Path(template_path).name}")
    print('='*60)

    print(f"\n총 문단 수: {result.total_paragraphs}")
    print(f"총 플레이스홀더 수: {len(result.placeholders)}")

    if result.placeholders:
        print("\n발견된 플레이스홀더:")
        for p in result.placeholders:
            print(f"  - {p.id}")
            print(f"    타입: {p.placeholder_type}")
            print(f"    섹션: {p.section_type}")
            print(f"    스타일: {p.style_name or 'None'}")
            print(f"    위치: 문단 {p.paragraph_index}")
            print()
    else:
        print("\n플레이스홀더가 발견되지 않았습니다.")
        print("템플릿에 {{TITLE}}, {{BODY}} 등의 플레이스홀더를 추가해주세요.")

    # JSON 출력
    print("\n--- JSON 출력 (LLM 입력용) ---")
    print(json.dumps(parser.to_dict(), ensure_ascii=False, indent=2))
