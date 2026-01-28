"""
DOCX 조립기 (Composer)

템플릿의 플레이스홀더를 실제 콘텐츠로 교체하여 최종 DOCX 생성
- 템플릿 복제
- 플레이스홀더 교체
- 스타일 보존
"""

import re
import shutil
from pathlib import Path
from typing import List, Optional, Dict, Any
from copy import deepcopy

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsmap
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from .models import (
    ContentMappingPlan, ContentMapping, ParsedTemplate,
    Placeholder, PlaceholderType, PLACEHOLDER_PATTERNS
)
from .markdown_parser import ContentBlock, DocumentStructure
from .template_analyzer import DocxTemplateAnalyzer, TemplateStructure


class DocxComposer:
    """
    DOCX 조립기

    템플릿과 매핑 계획을 바탕으로 최종 DOCX 생성
    """

    def __init__(
        self,
        template_path: str,
        output_dir: Optional[str] = None,
    ):
        """
        Args:
            template_path: 템플릿 DOCX 파일 경로
            output_dir: 출력 디렉토리 (기본: 템플릿과 같은 폴더)
        """
        self.template_path = Path(template_path)
        self.output_dir = Path(output_dir) if output_dir else self.template_path.parent

        # 템플릿 분석
        self._analyzer = DocxTemplateAnalyzer(str(self.template_path))
        self.template_structure = self._analyzer.analyze()

        # 플레이스홀더 패턴
        self.placeholder_regex = re.compile(PLACEHOLDER_PATTERNS["default"])

    def compose(
        self,
        mapping_plan: ContentMappingPlan,
        content: DocumentStructure,
        output_filename: Optional[str] = None,
    ) -> str:
        """
        최종 DOCX 생성

        Args:
            mapping_plan: 콘텐츠 매핑 계획
            content: 파싱된 마크다운 콘텐츠
            output_filename: 출력 파일명 (기본: template_output.docx)

        Returns:
            생성된 파일 경로
        """
        # 출력 경로 결정
        if output_filename:
            output_path = self.output_dir / output_filename
        else:
            output_path = self.output_dir / f"{self.template_path.stem}_output.docx"

        # 템플릿 복사하여 열기
        shutil.copy(self.template_path, output_path)
        doc = Document(str(output_path))

        # 문단별로 플레이스홀더 교체
        self._replace_placeholders_in_document(doc, mapping_plan, content)

        # 헤더/푸터의 플레이스홀더도 교체
        for section in doc.sections:
            if section.header:
                self._replace_placeholders_in_container(
                    section.header, mapping_plan, content
                )
            if section.footer:
                self._replace_placeholders_in_container(
                    section.footer, mapping_plan, content
                )

        # 저장
        doc.save(str(output_path))
        return str(output_path)

    def _replace_placeholders_in_document(
        self,
        doc: Document,
        mapping_plan: ContentMappingPlan,
        content: DocumentStructure,
    ):
        """문서 본문의 플레이스홀더 교체"""
        paragraphs_to_process = list(doc.paragraphs)

        for para in paragraphs_to_process:
            self._process_paragraph(para, mapping_plan, content, doc)

    def _replace_placeholders_in_container(
        self,
        container,
        mapping_plan: ContentMappingPlan,
        content: DocumentStructure,
    ):
        """헤더/푸터 등 컨테이너의 플레이스홀더 교체"""
        for para in container.paragraphs:
            self._process_paragraph_simple(para, mapping_plan, content)

    def _process_paragraph(
        self,
        para: Paragraph,
        mapping_plan: ContentMappingPlan,
        content: DocumentStructure,
        doc: Document,
    ):
        """단일 문단 처리 (본문용 - 다중 블록 삽입 지원)"""
        text = para.text
        matches = list(self.placeholder_regex.finditer(text))

        if not matches:
            return

        # 각 플레이스홀더 처리
        for match in matches:
            placeholder_id = match.group(0)  # "{{TITLE}}"
            mapping = mapping_plan.get_mapping_for_placeholder(placeholder_id)

            if not mapping:
                continue

            # 매핑된 콘텐츠 블록들
            blocks = [content.raw_blocks[i] for i in mapping.content_block_indices
                      if i < len(content.raw_blocks)]

            if not blocks:
                # 빈 매핑 - 플레이스홀더만 제거
                self._replace_text_in_paragraph(para, placeholder_id, "")
                continue

            # 단일 블록이면 텍스트 교체
            if len(blocks) == 1:
                replacement_text = self._get_block_text(blocks[0])
                self._replace_text_in_paragraph(para, placeholder_id, replacement_text)
            else:
                # 다중 블록이면 첫 블록으로 교체 후 추가 블록 삽입
                first_block = blocks[0]
                self._replace_text_in_paragraph(para, placeholder_id, self._get_block_text(first_block))

                # 나머지 블록들을 문단 뒤에 추가
                for block in blocks[1:]:
                    self._insert_block_after(para, block, doc)

    def _process_paragraph_simple(
        self,
        para: Paragraph,
        mapping_plan: ContentMappingPlan,
        content: DocumentStructure,
    ):
        """단일 문단 처리 (헤더/푸터용 - 단순 텍스트 교체)"""
        text = para.text
        matches = list(self.placeholder_regex.finditer(text))

        if not matches:
            return

        for match in matches:
            placeholder_id = match.group(0)
            mapping = mapping_plan.get_mapping_for_placeholder(placeholder_id)

            if not mapping:
                continue

            blocks = [content.raw_blocks[i] for i in mapping.content_block_indices
                      if i < len(content.raw_blocks)]

            if blocks:
                # 모든 블록을 하나의 텍스트로 합침
                replacement_text = " | ".join(self._get_block_text(b) for b in blocks)
            else:
                replacement_text = ""

            self._replace_text_in_paragraph(para, placeholder_id, replacement_text)

    def _get_block_text(self, block: ContentBlock) -> str:
        """ContentBlock에서 텍스트 추출"""
        if block.block_type == "list":
            # 리스트는 자식 아이템들의 텍스트
            items = []
            for child in block.children:
                prefix = "- " if block.list_type == "bullet" else f"{len(items)+1}. "
                items.append(prefix + child.content)
            return "\n".join(items)
        elif block.block_type == "table":
            # 테이블은 간단한 텍스트 표현
            rows = block.attributes.get("rows", [])
            lines = []
            for row in rows:
                cells = row.get("cells", [])
                lines.append(" | ".join(cells))
            return "\n".join(lines)
        else:
            return block.content

    def _replace_text_in_paragraph(
        self,
        para: Paragraph,
        old_text: str,
        new_text: str,
    ):
        """문단 내 텍스트 교체 (스타일 유지)"""
        # run 단위로 처리해야 스타일이 유지됨
        for run in para.runs:
            if old_text in run.text:
                run.text = run.text.replace(old_text, new_text)

    def _insert_block_after(
        self,
        para: Paragraph,
        block: ContentBlock,
        doc: Document,
    ):
        """문단 뒤에 새 블록 삽입"""
        # python-docx는 insert_paragraph_before만 지원
        # 따라서 다음 문단 앞에 삽입하거나 끝에 추가

        # 문단의 XML 요소 찾기
        para_element = para._element

        # 새 문단 생성
        if block.block_type == "heading":
            # 헤딩 레벨에 따른 스타일 적용
            style_name = f"Heading {block.level}" if block.level <= 6 else None
            new_para = doc.add_paragraph(block.content, style=style_name)
        elif block.block_type == "list":
            # 리스트 아이템들 추가
            for child in block.children:
                prefix = "- " if block.list_type == "bullet" else ""
                new_para = doc.add_paragraph(prefix + child.content)
        elif block.block_type == "code":
            new_para = doc.add_paragraph()
            run = new_para.add_run(block.content)
            run.font.name = "Consolas"
            run.font.size = Pt(9)
        else:
            new_para = doc.add_paragraph(block.content)

    def compose_with_sections(
        self,
        mapping_plan: ContentMappingPlan,
        content: DocumentStructure,
        output_filename: Optional[str] = None,
    ) -> str:
        """
        섹션별 조립 (고급 버전)

        표지, 목차, 본문 등 섹션별로 다른 처리 적용
        """
        output_path = self.output_dir / (output_filename or f"{self.template_path.stem}_output.docx")

        shutil.copy(self.template_path, output_path)
        doc = Document(str(output_path))

        # 문서 분석 결과 활용
        page_structure = self.template_structure.page_structure

        if page_structure:
            # 페이지 구조가 분석되어 있으면 활용
            self._compose_by_page_structure(doc, mapping_plan, content, page_structure)
        else:
            # 기본 플레이스홀더 교체
            self._replace_placeholders_in_document(doc, mapping_plan, content)

        doc.save(str(output_path))
        return str(output_path)

    def _compose_by_page_structure(
        self,
        doc: Document,
        mapping_plan: ContentMappingPlan,
        content: DocumentStructure,
        page_structure,
    ):
        """페이지 구조 기반 조립"""
        # 표지 처리
        if page_structure.cover_page is not None:
            cover_info = page_structure.pages[page_structure.cover_page]
            self._process_cover_section(doc, mapping_plan, content, cover_info)

        # 본문 처리
        self._replace_placeholders_in_document(doc, mapping_plan, content)

    def _process_cover_section(
        self,
        doc: Document,
        mapping_plan: ContentMappingPlan,
        content: DocumentStructure,
        cover_info,
    ):
        """표지 섹션 처리"""
        # TITLE, SUBTITLE 플레이스홀더 우선 처리
        for para in doc.paragraphs[:10]:  # 처음 10개 문단만 체크 (표지 영역)
            self._process_paragraph_simple(para, mapping_plan, content)


def compose_document(
    template_path: str,
    markdown_path: str,
    output_path: Optional[str] = None,
    use_llm: bool = False,
) -> str:
    """
    헬퍼 함수: 마크다운 + 템플릿 → DOCX

    Args:
        template_path: 템플릿 파일 경로
        markdown_path: 마크다운 파일 경로
        output_path: 출력 파일 경로
        use_llm: LLM 사용 여부

    Returns:
        생성된 파일 경로
    """
    from .template_parser import TemplateParser
    from .markdown_parser import MarkdownParser
    from .llm_content_mapper import LLMContentMapper, ContentMapperSync

    # 1. 템플릿 파싱
    template_parser = TemplateParser(template_path)
    template = template_parser.parse()

    # 2. 마크다운 파싱
    md_parser = MarkdownParser()
    content = md_parser.parse_file(markdown_path)

    # 3. 매핑 생성
    if use_llm:
        import asyncio
        async def get_mapping():
            async with LLMContentMapper(use_llm=True) as mapper:
                return await mapper.create_mapping_plan(template, content)
        mapping_plan = asyncio.run(get_mapping())
    else:
        mapper = ContentMapperSync()
        mapping_plan = mapper.create_mapping_plan(template, content)

    # 4. 문서 조립
    composer = DocxComposer(template_path)
    return composer.compose(
        mapping_plan=mapping_plan,
        content=content,
        output_filename=Path(output_path).name if output_path else None,
    )


if __name__ == "__main__":
    import sys

    if len(sys.argv) < 3:
        print("Usage: python -m src.docx_composer <template.docx> <content.md> [output.docx]")
        sys.exit(1)

    template = sys.argv[1]
    markdown = sys.argv[2]
    output = sys.argv[3] if len(sys.argv) > 3 else None

    result = compose_document(template, markdown, output)
    print(f"Generated: {result}")
