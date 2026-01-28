from docx import Document
from docx.shared import Pt
import re

def convert_md_to_custom_word(md_file_path, template_path, output_path):
    # 1. 템플릿 파일 불러오기
    doc = Document(template_path)

    # 2. 기존 템플릿의 본문 내용 비우기 (스타일 설정은 유지됨)
    #    단, '제목(Title)'이 있는 첫 페이지나 특정 섹션은 남겨두고 본문만 지울 수도 있으나,
    #    여기서는 깔끔하게 바디를 초기화하고 다시 채우는 방식을 예시로 듭니다.
    for element in doc.element.body:
        doc.element.body.remove(element)

    # 3. 마크다운 파일 읽기
    with open(md_file_path, 'r', encoding='utf-8') as f:
        md_lines = f.readlines()

    # 4. 한 줄씩 읽어서 워드 스타일 적용
    for line in md_lines:
        line = line.strip()
        if not line:
            continue

        # (1) 문서 제목 처리 (가장 첫 번째 큰 제목을 문서 제목으로 간주)
        # 템플릿 규칙상 Title 스타일을 쓰면 푸터도 자동 업데이트될 수 있음
        if line.startswith('**') and line.endswith('**') and '계약서' in line: 
             # 예: **부동산 매매계약서 / 토지철거 및 이설 계약**
            clean_text = line.replace('**', '')
            # 워드 템플릿의 'Title' 스타일 적용 (푸터 연동용)
            doc.add_paragraph(clean_text, style='Title')
            
        # (2) 헤딩 1 (# ) -> 템플릿의 'Heading 1'
        elif line.startswith('# ') or (line[0].isdigit() and '. **' in line):
            # MD 예시: "1. **계약서 정리**" 같은 패턴도 헤딩으로 처리
            clean_text = re.sub(r'[\*\#]', '', line).strip()
            doc.add_paragraph(clean_text, style='Heading 1')

        # (3) 헤딩 2 (## )
        elif line.startswith('## '):
            clean_text = line.replace('## ', '')
            doc.add_paragraph(clean_text, style='Heading 2')

        # (4) 리스트 (Bullet) -> 템플릿의 'List Bullet'
        elif line.startswith('- '):
            clean_text = line.replace('- ', '')
            doc.add_paragraph(clean_text, style='List Bullet')

        # (5) 번호 리스트 (1. ) -> 템플릿의 'List Number'
        elif re.match(r'^\d+\.', line):
            # "1. ", "2. " 등의 패턴 제거 혹은 유지 선택
            # 워드 스타일이 자동으로 번호를 매겨준다면 텍스트만 넣음
            clean_text = re.sub(r'^\d+\.\s*', '', line) 
            doc.add_paragraph(clean_text, style='List Number')

        # (6) 일반 본문
        else:
            # 볼드체(**) 처리 등 인라인 스타일링이 필요하면 여기서 추가 처리
            # 단순 텍스트 삽입 예시:
            p = doc.add_paragraph(style='Body Text') # 템플릿의 Body Text 스타일 사용
            # 부분 볼드 처리 로직 (간소화됨)
            if '**' in line:
                parts = line.split('**')
                for i, part in enumerate(parts):
                    run = p.add_run(part)
                    if i % 2 == 1: # 홀수 번째 조각은 ** 사이의 텍스트
                        run.bold = True
            else:
                p.add_run(line)

    # 5. 문서 속성 업데이트 (푸터가 문서 속성을 참조하는 경우)
    doc.core_properties.title = "부동산 매매계약서"
    
    # 6. 저장
    doc.save(output_path)
    print(f"변환 완료: {output_path}")

# 실제 사용 시
convert_md_to_custom_word('/home/shaush/work/parsed-outputs/예상산출물_한주 토지매각 관련 계약서_IRAC 2.md', '/home/shaush/md-to-docx/docx_only/[Word 템플릿] A4.docx', '최종결과물.docx')