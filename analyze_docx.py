"""
DOCX 템플릿 분석기
- 스타일(폰트, 크기, 색상 등) 추출
- 배경이미지(표지/본문) 추출
- 헤더/푸터 정보 추출
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
import zipfile
import os
import json
from pathlib import Path


def extract_font_info(font):
    """폰트 정보 추출"""
    info = {}
    if font.name:
        info['name'] = font.name
    if font.size:
        info['size_pt'] = font.size.pt
    if font.bold is not None:
        info['bold'] = font.bold
    if font.italic is not None:
        info['italic'] = font.italic
    if font.underline is not None:
        info['underline'] = font.underline
    if font.color and font.color.rgb:
        info['color'] = str(font.color.rgb)
    return info


def extract_paragraph_format(pf):
    """문단 포맷 정보 추출"""
    info = {}
    if pf.alignment:
        info['alignment'] = str(pf.alignment)
    if pf.line_spacing:
        info['line_spacing'] = pf.line_spacing
    if pf.space_before:
        info['space_before_pt'] = pf.space_before.pt
    if pf.space_after:
        info['space_after_pt'] = pf.space_after.pt
    if pf.left_indent:
        info['left_indent_pt'] = pf.left_indent.pt
    if pf.first_line_indent:
        info['first_line_indent_pt'] = pf.first_line_indent.pt
    return info


def analyze_styles(doc):
    """문서의 모든 스타일 분석"""
    styles_info = {}

    for style in doc.styles:
        if style.type == WD_STYLE_TYPE.PARAGRAPH:
            style_data = {
                'type': 'paragraph',
                'name': style.name,
                'base_style': style.base_style.name if style.base_style else None,
            }

            # 폰트 정보
            if style.font:
                style_data['font'] = extract_font_info(style.font)

            # 문단 포맷
            if style.paragraph_format:
                style_data['paragraph_format'] = extract_paragraph_format(style.paragraph_format)

            styles_info[style.name] = style_data

        elif style.type == WD_STYLE_TYPE.CHARACTER:
            style_data = {
                'type': 'character',
                'name': style.name,
            }
            if style.font:
                style_data['font'] = extract_font_info(style.font)
            styles_info[style.name] = style_data

    return styles_info


def analyze_document_content(doc):
    """문서 본문 내용 분석 (실제 사용된 스타일 파악)"""
    content_analysis = []

    for i, para in enumerate(doc.paragraphs):
        para_info = {
            'index': i,
            'text_preview': para.text[:100] if para.text else '',
            'style_name': para.style.name if para.style else None,
        }

        # 실제 적용된 폰트 정보 (run 레벨)
        runs_info = []
        for run in para.runs:
            run_info = {
                'text_preview': run.text[:50] if run.text else '',
            }
            if run.font:
                run_info['font'] = extract_font_info(run.font)
            runs_info.append(run_info)

        if runs_info:
            para_info['runs'] = runs_info[:3]  # 처음 3개만

        content_analysis.append(para_info)

    return content_analysis[:20]  # 처음 20개 문단만


def extract_images_from_docx(docx_path, output_dir):
    """DOCX에서 모든 이미지 추출 (배경 포함)"""
    images = []
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    with zipfile.ZipFile(docx_path, 'r') as zip_ref:
        for file_name in zip_ref.namelist():
            # 이미지 파일 찾기 (word/media/ 폴더)
            if file_name.startswith('word/media/'):
                image_name = os.path.basename(file_name)
                image_path = output_dir / image_name

                with zip_ref.open(file_name) as src:
                    with open(image_path, 'wb') as dst:
                        dst.write(src.read())

                images.append({
                    'original_path': file_name,
                    'extracted_path': str(image_path),
                    'name': image_name
                })

            # 배경 관련 파일 (theme, settings 등)
            if 'theme' in file_name.lower() or 'background' in file_name.lower():
                print(f"Found theme/background related: {file_name}")

    return images


def analyze_headers_footers(doc):
    """헤더/푸터 분석"""
    headers_footers = {'headers': [], 'footers': []}

    for section in doc.sections:
        # 헤더
        header = section.header
        if header:
            header_text = '\n'.join([p.text for p in header.paragraphs])
            headers_footers['headers'].append({
                'text': header_text,
                'paragraphs_count': len(header.paragraphs)
            })

        # 푸터
        footer = section.footer
        if footer:
            footer_text = '\n'.join([p.text for p in footer.paragraphs])
            headers_footers['footers'].append({
                'text': footer_text,
                'paragraphs_count': len(footer.paragraphs)
            })

    return headers_footers


def analyze_document_properties(doc):
    """문서 속성 분석"""
    props = doc.core_properties
    return {
        'title': props.title,
        'author': props.author,
        'subject': props.subject,
        'keywords': props.keywords,
        'created': str(props.created) if props.created else None,
        'modified': str(props.modified) if props.modified else None,
    }


def analyze_sections(doc):
    """섹션 정보 분석 (페이지 설정)"""
    sections_info = []
    for i, section in enumerate(doc.sections):
        sec_info = {
            'index': i,
            'page_width_inches': section.page_width.inches if section.page_width else None,
            'page_height_inches': section.page_height.inches if section.page_height else None,
            'left_margin_inches': section.left_margin.inches if section.left_margin else None,
            'right_margin_inches': section.right_margin.inches if section.right_margin else None,
            'top_margin_inches': section.top_margin.inches if section.top_margin else None,
            'bottom_margin_inches': section.bottom_margin.inches if section.bottom_margin else None,
        }
        sections_info.append(sec_info)
    return sections_info


def extract_background_from_xml(docx_path):
    """XML에서 배경 정보 직접 추출"""
    backgrounds = []

    with zipfile.ZipFile(docx_path, 'r') as zip_ref:
        # document.xml에서 배경 정보 찾기
        if 'word/document.xml' in zip_ref.namelist():
            with zip_ref.open('word/document.xml') as f:
                content = f.read().decode('utf-8')
                # 배경 관련 태그 검색
                if 'background' in content.lower():
                    backgrounds.append({'type': 'document_background', 'found': True})
                if 'w:bg' in content or 'a:blip' in content:
                    backgrounds.append({'type': 'image_background', 'found': True})

        # settings.xml 확인
        if 'word/settings.xml' in zip_ref.namelist():
            with zip_ref.open('word/settings.xml') as f:
                content = f.read().decode('utf-8')
                if 'displayBackgroundShape' in content:
                    backgrounds.append({'type': 'background_shape_enabled', 'found': True})

        # _rels 파일 목록
        rels_files = [f for f in zip_ref.namelist() if f.endswith('.rels')]
        backgrounds.append({'rels_files': rels_files})

    return backgrounds


def full_analysis(docx_path, output_dir=None):
    """전체 분석 실행"""
    print(f"\n{'='*60}")
    print(f"Analyzing: {docx_path}")
    print('='*60)

    doc = Document(docx_path)

    if output_dir is None:
        output_dir = Path(docx_path).stem + '_extracted'

    analysis = {
        'file': str(docx_path),
        'properties': analyze_document_properties(doc),
        'sections': analyze_sections(doc),
        'styles': analyze_styles(doc),
        'content_preview': analyze_document_content(doc),
        'headers_footers': analyze_headers_footers(doc),
        'backgrounds': extract_background_from_xml(docx_path),
    }

    # 이미지 추출
    images = extract_images_from_docx(docx_path, output_dir)
    analysis['images'] = images

    return analysis


def print_key_styles(analysis):
    """주요 스타일 출력"""
    print("\n📌 주요 스타일 정보:")
    key_styles = ['Title', 'Heading 1', 'Heading 2', 'Heading 3',
                  'Normal', 'Body Text', 'List Bullet', 'List Number']

    for style_name in key_styles:
        if style_name in analysis['styles']:
            style = analysis['styles'][style_name]
            print(f"\n  [{style_name}]")
            if 'font' in style:
                font = style['font']
                print(f"    폰트: {font.get('name', 'N/A')}, 크기: {font.get('size_pt', 'N/A')}pt")
                print(f"    Bold: {font.get('bold', 'N/A')}, Italic: {font.get('italic', 'N/A')}")


if __name__ == '__main__':
    import sys

    # 기본 템플릿 분석
    template_path = '/home/shaush/md-to-docx/docx_only/[Word 템플릿] A4.docx'

    if len(sys.argv) > 1:
        template_path = sys.argv[1]

    analysis = full_analysis(template_path, 'extracted_template')

    # 주요 정보 출력
    print("\n📄 문서 속성:")
    for k, v in analysis['properties'].items():
        if v:
            print(f"  {k}: {v}")

    print("\n📐 페이지 설정:")
    for sec in analysis['sections']:
        print(f"  섹션 {sec['index']}: {sec['page_width_inches']:.2f}\" x {sec['page_height_inches']:.2f}\"")
        print(f"    여백: L={sec['left_margin_inches']:.2f}\", R={sec['right_margin_inches']:.2f}\", T={sec['top_margin_inches']:.2f}\", B={sec['bottom_margin_inches']:.2f}\"")

    print_key_styles(analysis)

    print("\n🖼️ 추출된 이미지:")
    for img in analysis['images']:
        print(f"  - {img['name']} ({img['original_path']})")

    print("\n📑 헤더/푸터:")
    hf = analysis['headers_footers']
    for i, h in enumerate(hf['headers']):
        print(f"  헤더 {i}: {h['text'][:50] if h['text'] else '(empty)'}...")
    for i, f in enumerate(hf['footers']):
        print(f"  푸터 {i}: {f['text'][:50] if f['text'] else '(empty)'}...")

    print("\n🎨 배경 정보:")
    for bg in analysis['backgrounds']:
        print(f"  {bg}")

    # JSON으로 저장
    output_json = 'template_analysis.json'
    with open(output_json, 'w', encoding='utf-8') as f:
        # datetime 등 직렬화 불가 객체 처리
        json.dump(analysis, f, ensure_ascii=False, indent=2, default=str)
    print(f"\n✅ 분석 결과 저장: {output_json}")
